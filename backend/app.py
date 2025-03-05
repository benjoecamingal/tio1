from flask import Flask, request, jsonify, session, send_file
from flask_cors import CORS
import pandas as pd
import os
import io
from datetime import datetime, timedelta
import openpyxl
from openpyxl.styles import NamedStyle
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from barcode import Code128
from barcode.writer import ImageWriter
import tempfile
import random
import string
from db_config import get_db_connection

app = Flask(__name__)
app.secret_key = "your_secret_key"
CORS(app, supports_credentials=True)

UPLOAD_FOLDER = tempfile.gettempdir()
PROCESSED_FOLDER = tempfile.gettempdir()

def generate_barcode_image(barcode_number, output_path):
    try:
        if output_path.endswith('.png'):
            output_path = output_path[:-4]
        barcode = Code128(str(barcode_number), writer=ImageWriter())
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        saved_path = barcode.save(output_path)
        return True
    except Exception as e:
        return False

def generate_word_document(dataframe, output_path, barcode_paths):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    for index, row in dataframe.iterrows():
        name = row.get("Name", "Unknown")
        barcode_number = row.get("Barcode", "")

        if not barcode_number:
            continue

        barcode_path = barcode_paths.get(barcode_number)

        if not barcode_path or not os.path.exists(barcode_path):
            continue

        # Use a single paragraph for both name and barcode image
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add name (Centered, Large, Bold) in the same paragraph
        run = paragraph.add_run(f"{name}\n")
        run.bold = True
        run.font.size = Pt(14)

        # Add barcode image in the same paragraph, inline with the name
        try:
            run_img = paragraph.add_run()
            run_img.add_picture(barcode_path, width=Inches(1.5), height=Inches(0.75))
        except Exception as e:
            error_p = doc.add_paragraph("[ERROR: Could not add barcode image]")
            error_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # doc.add_paragraph("-" * 30).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    try:
        doc.save(output_path)
    except Exception as e:
        pass

def generate_unique_barcode(cursor):
    while True:
        barcode = ''.join(random.choices(string.digits, k=12))  # Simple barcode generation
        cursor.execute("SELECT COUNT(*) FROM students WHERE barcode = %s", (barcode,))
        if cursor.fetchone()[0] == 0:
            return barcode

def cleanup_files(temp_file_path, output_docx_path, barcode_paths):
    if temp_file_path and os.path.exists(temp_file_path):
        os.unlink(temp_file_path)
    if output_docx_path and os.path.exists(output_docx_path):
        os.unlink(output_docx_path)
    for path in barcode_paths.values():
        if os.path.exists(path):
            os.unlink(path)

def format_excel_time(worksheet):
    time_style = NamedStyle(name='time_style', number_format='HH:MM:SS')
    headers = [cell.value for cell in worksheet[1]]
    time_in_col = headers.index('Time In') + 1 if 'Time In' in headers else None
    time_out_col = headers.index('Time Out') + 1 if 'Time Out' in headers else None
    
    if time_in_col:
        for cell in worksheet[get_column_letter(time_in_col)][1:]:
            cell.style = time_style
    if time_out_col:
        for cell in worksheet[get_column_letter(time_out_col)][1:]:
            cell.style = time_style

# Routes
@app.route("/login", methods=["POST"])
def login():
    data = request.json
    username = data.get("username")
    password = data.get("password")

    db = get_db_connection()
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE username = %s AND password = %s", (username, password))
    user = cursor.fetchone()
    db.close()

    if user:
        session["user"] = username
        session.permanent = True
        return jsonify({"message": "Login successful"}), 200
    return jsonify({"error": "Invalid credentials"}), 401

@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
    file.save(temp_file.name)
    temp_file.close()

    try:
        df = pd.read_excel(temp_file.name, engine="openpyxl")
        df.dropna(how="all", inplace=True)
        df.fillna("", inplace=True)

        required_columns = ["Name", "Batch", "Position", "Department", "School"]
        if not all(col in df.columns for col in required_columns):
            raise ValueError("Missing required columns")

        db = get_db_connection()
        cursor = db.cursor()

        if "Barcode" not in df.columns:
            df["Barcode"] = df.apply(lambda _: generate_unique_barcode(cursor), axis=1)

        barcode_paths = {}
        for _, row in df.iterrows():
            barcode_number = row["Barcode"]
            base_path = os.path.join(tempfile.gettempdir(), f"barcode_{barcode_number}")
            barcode_path = f"{base_path}.png"
            if generate_barcode_image(barcode_number, base_path):
                barcode_paths[barcode_number] = barcode_path

        for _, row in df.iterrows():
            cursor.execute(
                "INSERT INTO students (Name, Batch, Position, Department, School, Barcode) VALUES (%s, %s, %s, %s, %s, %s)",
                (row["Name"], row["Batch"], row["Position"], row["Department"], row["School"], row["Barcode"])
            )

        db.commit()
        output_docx = os.path.join(tempfile.gettempdir(), "student_barcodes.docx")
        generate_word_document(df, output_docx, barcode_paths)

        with open(output_docx, 'rb') as f:
            file_data = io.BytesIO(f.read())

        cleanup_files(temp_file.name, output_docx, barcode_paths)
        file_data.seek(0)
        return send_file(
            file_data,
            as_attachment=True,
            download_name="student_barcodes.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        cleanup_files(temp_file.name, None, {})
        return jsonify({"error": str(e)}), 500

@app.route("/add_student", methods=["POST"])
def add_student():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 403

    data = request.get_json()
    required_fields = ["name", "batch", "position", "department", "school"]
    if not all(data.get(field) for field in required_fields):
        return jsonify({"error": "Missing required fields"}), 400

    df = pd.DataFrame([{
        'Name': data["name"].strip(),
        'Batch': data["batch"].strip(),
        'Position': data["position"].strip(),
        'Department': data["department"].strip(),
        'School': data["school"].strip()
    }])

    try:
        db = get_db_connection()
        cursor = db.cursor()
        barcode = generate_unique_barcode(cursor)
        df["Barcode"] = barcode

        base_path = os.path.join(tempfile.gettempdir(), f"barcode_{barcode}")
        barcode_path = f"{base_path}.png"
        if generate_barcode_image(barcode, base_path):
            barcode_paths = {barcode: barcode_path}
        else:
            raise Exception("Failed to generate barcode")

        cursor.execute(
            "INSERT INTO students (Name, Batch, Position, Department, School, Barcode) VALUES (%s, %s, %s, %s, %s, %s)",
            (df.iloc[0]["Name"], df.iloc[0]["Batch"], df.iloc[0]["Position"], 
             df.iloc[0]["Department"], df.iloc[0]["School"], barcode)
        )

        db.commit()
        output_docx = os.path.join(tempfile.gettempdir(), f"student_barcode_{barcode}.docx")
        generate_word_document(df, output_docx, barcode_paths)

        with open(output_docx, 'rb') as f:
            file_data = io.BytesIO(f.read())

        cleanup_files(None, output_docx, barcode_paths)
        file_data.seek(0)
        return send_file(
            file_data,
            as_attachment=True,
            download_name=f"student_barcode_{data['name']}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        db.rollback()
        cleanup_files(None, None, barcode_paths)
        return jsonify({"error": str(e)}), 500

@app.route("/logout", methods=["POST"])
def logout():
    session.pop("user", None)
    return jsonify({"message": "Logged out successfully"}), 200

@app.route('/scan', methods=['POST'])
def process_scan():
    data = request.get_json()
    barcode = data.get("barcode")

    if not barcode:
        return jsonify({"success": False, "message": "No barcode received"}), 400

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, department FROM students WHERE barcode = %s", (barcode.strip(),))
        student = cursor.fetchone()

        if not student:
            return jsonify({"success": False, "message": "Student not found"}), 404

        student_id, student_name, department = student
        cursor.execute(
            "SELECT id, time_in, time_out FROM attendance WHERE student_id = %s AND date = CURDATE()",
            (student_id,))
        record = cursor.fetchone()

        if record:
            attendance_id, time_in, time_out = record
            if time_in and not time_out:
                cursor.execute(
                    "UPDATE attendance SET time_out = CURRENT_TIME, date = CURDATE() WHERE id = %s",
                    (attendance_id,))
                conn.commit()
                time_out_formatted = datetime.now().strftime("%I:%M %p") if time_out is None else datetime.strptime(str(time_out), "%H:%M:%S").strftime("%I:%M %p")
                time_in_formatted = datetime.strptime(str(time_in), "%H:%M:%S").strftime("%I:%M %p") if time_in else "N/A"
                return jsonify({
                    "success": True,
                    "message": f"Time Out recorded for {student_name}",
                    "name": student_name,
                    "department": department,
                    "time_in": time_in_formatted,
                    "time_out": time_out_formatted,
                    "status": "Time Out",
                    "date": "Today"
                })
            time_in_formatted = datetime.strptime(str(time_in), "%H:%M:%S").strftime("%I:%M %p") if time_in else "N/A"
            time_out_formatted = datetime.strptime(str(time_out), "%H:%M:%S").strftime("%I:%M %p") if time_out else "N/A"
            return jsonify({
                "success": False,
                "message": "Already Timed Out for Today",
                "name": student_name,
                "department": department,
                "time_in": time_in_formatted,
                "time_out": time_out_formatted,
                "status": "Already Timed Out",
                "date": "Today"
            })
        else:
            current_time = datetime.now()
            time_in_formatted = current_time.strftime("%I:%M %p")
            cursor.execute(
                "INSERT INTO attendance (student_id, date, time_in) VALUES (%s, CURDATE(), %s)",
                (student_id, current_time.time()))
            conn.commit()
            return jsonify({
                "success": True,
                "message": f"Time In recorded for {student_name}",
                "name": student_name,
                "department": department,
                "time_in": time_in_formatted,
                "time_out": "N/A",
                "status": "Time In",
                "date": "Today"
            })

    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        conn.close()

@app.route('/attendance', methods=['GET'])
def get_attendance():
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        batch = request.args.get('batch')
        position = request.args.get('position')
        department = request.args.get('department')
        school = request.args.get('school')
        date = request.args.get('date')

        query = """
            SELECT s.name, s.batch, s.position, s.department, s.school, 
                   a.date, a.time_in, a.time_out
            FROM attendance a
            JOIN students s ON a.student_id = s.id
            WHERE 1=1
        """
        filters = []

        if batch:
            query += " AND s.batch = %s"
            filters.append(batch)
        if position:
            query += " AND s.position = %s"
            filters.append(position)
        if department:
            query += " AND s.department = %s"
            filters.append(department)
        if school:
            query += " AND s.school = %s"
            filters.append(school)
        if date:
            query += " AND a.date = %s"
            filters.append(date)

        query += " ORDER BY a.date DESC"

        cursor.execute(query, tuple(filters))
        records = cursor.fetchall()

        for record in records:
            record["time_in"] = str(record["time_in"]) if record["time_in"] else "N/A"
            record["time_out"] = str(record["time_out"]) if record["time_out"] else "N/A"

        return jsonify(records)

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()  # Ensure database connection is closed

@app.route('/filters', methods=['GET'])
def get_filters():
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)

        cursor.execute("SELECT DISTINCT batch FROM students WHERE batch IS NOT NULL")
        batches = [row["batch"] for row in cursor.fetchall()]

        cursor.execute("SELECT DISTINCT position FROM students WHERE position IS NOT NULL")
        positions = [row["position"] for row in cursor.fetchall()]

        cursor.execute("SELECT DISTINCT department FROM students WHERE department IS NOT NULL")
        departments = [row["department"] for row in cursor.fetchall()]

        cursor.execute("SELECT DISTINCT school FROM students WHERE school IS NOT NULL")
        schools = [row["school"] for row in cursor.fetchall()]

        return jsonify({
            "batches": batches,
            "positions": positions,
            "departments": departments,
            "schools": schools
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()  # Ensure database connection is closed

@app.route('/attendance/download', methods=['GET'])
def download_attendance():
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        batch = request.args.get('batch')
        position = request.args.get('position')
        department = request.args.get('department')
        school = request.args.get('school')
        date = request.args.get('date')

        query = """
            SELECT 
                s.name as 'Name',
                s.batch as 'Batch',
                s.position as 'Position',
                s.department as 'Department',
                s.school as 'School',
                a.date as 'Date',
                a.time_in as 'Time In',
                a.time_out as 'Time Out'
            FROM attendance a
            JOIN students s ON a.student_id = s.id
            WHERE 1=1
        """
        filters = []

        if batch:
            query += " AND s.batch = %s"
            filters.append(batch)
        if position:
            query += " AND s.position = %s"
            filters.append(position)
        if department:
            query += " AND s.department = %s"
            filters.append(department)
        if school:
            query += " AND s.school = %s"
            filters.append(school)
        if date:
            query += " AND a.date = %s"
            filters.append(date)

        query += " ORDER BY a.date DESC, s.name ASC"

        cursor.execute(query, tuple(filters))
        records = cursor.fetchall()

        total_students = len(set(record['Name'] for record in records))

        for record in records:
            if isinstance(record['Time In'], timedelta):
                record['Time In'] = str(record['Time In'])
            if isinstance(record['Time Out'], timedelta):
                record['Time Out'] = str(record['Time Out'])

        df = pd.DataFrame(records)
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Attendance', index=False)
            worksheet = writer.sheets['Attendance']
            worksheet.append(["Total Students:", total_students])
            
            for idx, col in enumerate(df.columns):
                max_length = max(df[col].astype(str).apply(len).max(), len(col)) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = max_length

            format_excel_time(worksheet)

        output.seek(0)
        current_date = datetime.now().strftime('%Y%m%d')
        filename = f"attendance_report_{current_date}.xlsx"

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()  # Ensure database connection is closed

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000, debug=True)